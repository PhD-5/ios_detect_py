#coding=utf-8
from docx import Document
import data
class Generator:

    def __init__(self):
        self.document = Document()

    def generate(self):
        self.write_head()
        self.write_app_info()
        self.write_binary_info()
        self.write_transport_info()
        self.write_storage_info()


        self.document.save('./temp/{}/report/{}.docx'.format(data.start_time, data.app_bundleID))

        #write app info

    def write_head(self):
        self.document.add_heading(u"iOS应用漏洞检测报告", level=0)


    def write_app_info(self):
        self.document.add_heading(u"应用详情",level=1)

        self.document.add_paragraph(u'检测时间：', style='List Bullet')
        self.document.add_paragraph(data.start_time)
        self.document.add_paragraph(u'应用ID：', style='List Bullet')
        self.document.add_paragraph(data.app_bundleID)
        self.document.add_paragraph(u'应用名称：', style='List Bullet')
        self.document.add_paragraph(data.metadata['name'])
        self.document.add_paragraph(u'应用版本：', style='List Bullet')
        self.document.add_paragraph(data.metadata['app_version'])
        self.document.add_paragraph(u'Bundle路径：', style='List Bullet')
        self.document.add_paragraph(data.metadata['bundle_directory'])
        self.document.add_paragraph(u'Data路径：', style='List Bullet')
        self.document.add_paragraph(data.metadata['data_directory'])

    def write_binary_info(self):
        self.document.add_heading(u"二进制检测结果：", level=1)

        self.document.add_heading(u"二进制保护措施检测", level=2)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        head_cell = table.rows[0].cells
        head_cell[0].text = u'架构'
        head_cell[1].text = 'Encrypted'
        head_cell[2].text = 'Stack Canaries'
        head_cell[3].text = 'ARC'
        head_cell[4].text = 'PIE'
        for arch in data.protection_check_lables.keys():
            row_cells = table.add_row().cells
            row_cells[0].text = arch
            row_cells[1].text = str(data.protection_check_lables[arch]['Encrypted'])
            row_cells[2].text = str(data.protection_check_lables[arch]['Stack Canaries'])
            row_cells[3].text = str(data.protection_check_lables[arch]['ARC'])
            row_cells[4].text = str(data.protection_check_lables[arch]['PIE'])

        self.document.add_heading(u"依赖库检测", level=2)
        table_share_lib = self.document.add_table(rows=1, cols=1)
        table_share_lib.style = 'Table Grid'
        head_cell_lib = table_share_lib.rows[0].cells
        head_cell_lib[0].text = 'Shared Library'
        for lib in data.shared_lib:
            row_lib_cells = table_share_lib.add_row().cells
            row_lib_cells[0].text = lib

        self.document.add_heading(u"硬编码", level=2)
        table_hard_code = self.document.add_table(rows=1, cols=1)
        table_hard_code.style = 'Table Grid'
        head_hard_code = table_hard_code.rows[0].cells
        head_hard_code[0].text = 'HardCode:'
        if len(data.hardcode) == 0:
            row_lib_cells = table_hard_code.add_row().cells
            row_lib_cells[0].text = u'无'
        for code in data.hardcode:
            row_lib_cells = table_hard_code.add_row().cells
            row_lib_cells[0].text = code

        self.document.add_heading(u"URL Fuzz", level=2)
        table_fuzz = self.document.add_table(rows=1, cols=2)
        table_fuzz.style = 'Table Grid'
        head_fuzz = table_fuzz.rows[0].cells
        head_fuzz[0].text = 'URL'
        head_fuzz[1].text = 'Result'
        for key in data.fuzz_result:
            row_cells = table_fuzz.add_row().cells
            row_cells[0].text = key
            if data.fuzz_result[key]:
                row_cells[1].text = 'Crash'
            else:
                row_cells[1].text = 'Safe'

    def write_transport_info(self):
        self.document.add_heading(u"传输层检测结果", level=1)
        self.document.add_heading(u"中间人攻击检测", level=2)
        cols_count = len(data.mitm_results.keys())+1
        table_1 = self.document.add_table(rows=1, cols=cols_count)
        table_1.style = 'Table Grid'
        head_1_cell = table_1.rows[0].cells
        head_1_cell[0].text = u'漏洞类型'
        i=1
        for item in data.mitm_results.keys():
            head_1_cell[i].text = data.mitm_results.keys()[i-1]
            i+=1
        row_1_cell = table_1.add_row().cells
        row_1_cell[0].text = u'漏洞统计'
        i=1
        for key in data.mitm_results.keys():
            row_1_cell[i].text = str(data.mitm_results[key])
            i+=1

        self.document.add_heading(u"不安全协议使用", level=2)
        table_2 = self.document.add_table(rows=1,cols=1)
        table_2.style = 'Table Grid'
        head_2_cell = table_2.rows[0].cells
        head_2_cell[0].text = u'未使用https'
        for item in data.traffic_unsafe_result:
            row_2_cell = table_2.add_row().cells
            row_2_cell[0].text = item
        if len(data.traffic_unsafe_result)==0:
            row_2_cell = table_2.add_row().cells
            row_2_cell[0].text = u"无"

        self.document.add_heading(u"敏感信息传输", level=2)
        table_3 = self.document.add_table(rows=1, cols=3)
        table_3.style = 'Table Grid'
        head_3_cell = table_3.rows[0].cells
        head_3_cell[0].text = 'url'
        head_3_cell[1].text = 'body'
        head_3_cell[2].text = u'敏感内容'
        for item in data.dynamic_sensitive_json['traffic']:
            print item
            row_3_cell = table_3.add_row().cells
            row_3_cell[0].text = item[0]['url']
            if item[0].has_key('body'):
                row_3_cell[1].text = item[0]['body']
            else:
                row_3_cell[1].text = u"无"
            row_3_cell[2].text = '-'.join(item[1])

    def write_storage_info(self):
        self.document.add_heading(u"存储层检测结果", level=1)

        self.document.add_heading(u"动态检测结果", level=2)
        self.document.add_heading(u"KeyChain", level=3)
        table1 = self.document.add_table(rows=1,cols=3)
        table1.style = 'Table Grid'
        head1_cells = table1.rows[0].cells
        head1_cells[0].text = u'函数名称'
        head1_cells[1].text = u'函数参数'
        head1_cells[2].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['keychain']:
            row1_cells = table1.add_row().cells
            row1_cells[0].text = item[0]['function']
            if item[0].has_key('attributes'):
                if item[0]['attributes'].has_key('kSecValueData'):
                    row1_cells[1].text = item[0]['attributes']['kSecValueData']
            elif item[0].has_key('attributesToUpdate'):
                if item[0]['attributesToUpdate'].has_key('kSecValueData'):
                    row1_cells[1].text = item[0]['attributesToUpdate']['kSecValueData']
            row1_cells[2].text = '-'.join(item[1])

        self.document.add_heading(u"NSUserDefaults", level=3)
        table2 = self.document.add_table(rows=1, cols=4)
        table2.style = 'Table Grid'
        head2_cells = table2.rows[0].cells
        head2_cells[0].text = u'key'
        head2_cells[1].text = u'写入内容'
        head2_cells[2].text = u'数据类型'
        head2_cells[3].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['nsuserdefaults']:
            row1_cells = table2.add_row().cells
            row1_cells[0].text = item[0]['key']
            row1_cells[1].text = str(item[0]['content'])
            row1_cells[2].text = item[0]['sourceType']
            row1_cells[3].text = '-'.join(item[1])

        self.document.add_heading(u"Plist", level=3)
        table3 = self.document.add_table(rows=1, cols=4)
        table3.style = 'Table Grid'
        head3_cells = table3.rows[0].cells
        head3_cells[0].text = u'文件路径'
        head3_cells[1].text = u'写入内容'
        head3_cells[2].text = u'数据类型'
        head3_cells[3].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['plist']:
            row1_cells = table3.add_row().cells
            if item[0].has_key('filepath'):
                row1_cells[0].text = item[0]['filepath']
            elif item[0].has_key('url'):
                row1_cells[0].text = item[0]['url']
            row1_cells[1].text = str(item[0]['content'])
            row1_cells[2].text = item[0]['sourceType']
            row1_cells[3].text = '-'.join(item[1])


        self.document.add_heading(u"本地审计检测结果", level=2)
        self.document.add_heading(u"KeyChain", level=3)
        self.document.add_heading(u"数据库文件", level=3)
        for file_path in data.db_file_results.keys():
            self.document.add_paragraph(u'文件路径：'+file_path)
            table5 = self.document.add_table(rows=1,cols=3)
            table5.style = 'Table Grid'
            head5_cells = table5.rows[0].cells
            head5_cells[0].text = u'表名'
            head5_cells[1].text = u'行信息'
            head5_cells[2].text = u'敏感数据'
            for item in data.db_file_results[file_path]:
                row5_cells = table5.add_row().cells
                row5_cells[0].text = item[0]
                row5_cells[1].text = str(item[1])
                row5_cells[2].text = item[2]

        self.document.add_heading(u"Plist文件", level=3)
        for file_path in data.plist_file_results.keys():
            self.document.add_paragraph(u'文件路径：'+file_path)
            table6 = self.document.add_table(rows=1,cols=3)
            table6.style = 'Table Grid'
            head6_cells = table6.rows[0].cells
            head6_cells[0].text = u'key路径'
            head6_cells[1].text = u'单元信息'
            head6_cells[2].text = u'敏感数据'
            for item in data.plist_file_results[file_path]:
                row6_cells = table6.add_row().cells
                row6_cells[0].text = item[0]
                row6_cells[1].text = str(item[1])
                row6_cells[2].text = item[2]

        self.document.add_heading(u"本地文件保护检测", level=3)
        table7 = self.document.add_table(rows=1,cols=2)
        table7.style = 'Table Grid'
        head7_cells = table7.rows[0].cells
        head7_cells[0].text = u'文件路径'
        head7_cells[1].text = u'保护信息'
        for item in data.local_file_protection:
            row7_cells = table7.add_row().cells
            row7_cells[0].text = item[0]
            row7_cells[1].text = item[1]

if __name__=='__main__':
    data = dict()
    Generator().generate()